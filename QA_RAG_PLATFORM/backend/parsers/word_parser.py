import io
from docx import Document
from typing import List, Dict, Any


def parse(content: bytes, filename: str) -> List[Dict[str, Any]]:
    doc = Document(io.BytesIO(content))
    pages = []
    current_text: List[str] = []
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_text.append(text)
        # Group every 30 paragraphs as one logical "page"
        if len(current_text) >= 30:
            pages.append({
                "text": "\n".join(current_text),
                "metadata": {"filename": filename, "source": filename, "page": page_num},
                "page": page_num,
            })
            current_text = []
            page_num += 1

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            current_text.extend(rows)

    if current_text:
        pages.append({
            "text": "\n".join(current_text),
            "metadata": {"filename": filename, "source": filename, "page": page_num},
            "page": page_num,
        })

    return pages or [{"text": "", "metadata": {"filename": filename}, "page": 1}]
